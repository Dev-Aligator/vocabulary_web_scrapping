from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import os


def create_custom_style(document, style_name, font_size):
    style = document.styles.add_style(style_name, 1)  # 1 for paragraph style
    style.font.size = font_size
    return style


def convert_docx(data):
    document = Document()
    custom_style_p = create_custom_style(document, 'CustomStyle', Pt(14))
    heading = document.add_heading(data[0], level=0)
    heading.alignment = 1  # 1 for center alignment, 0 for left alignment, 2 for right alignment
    # p = document.add_paragraph('A plain paragraph having some ')
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True

    for section in data[1]:

        section_header = document.add_heading(section[0], level=1)
        for run in section_header.runs:
            run.font.size = Pt(18)
        for id, word in enumerate(section[1:]):
            p = document.add_paragraph(str(id + 1) + ". ", style='CustomStyle')
            p.add_run(word.get('word')).bold = True
            p.add_run(' ' + word.get('definition'))
            # Process english_example and bold the word
            p = document.add_paragraph("  ", style='CustomStyle')
            english_example = word.get('english_example')
            example_parts = english_example.split(' ')  # Split into words
            
            for i, part in enumerate(example_parts):
                if part.lower() == word.get('word').lower() or part.lower()[:-1] == word.get('word').lower():  # Compare word ignoring case
                    run = p.add_run(part + ' ')
                    run.bold = True
                else:
                    p.add_run(part + ' ')

            p.add_run(word.get('vietnamese_example'))

    output_filename = data[0].replace(" ", "_").replace(".", "").lower() + '.docx'
    output_path = os.path.join('output', output_filename)
    document.save(output_path)
